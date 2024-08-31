import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

country_names = []
percent_lists = []

try:
    url = 'https://en.wikipedia.org/wiki/List_of_Asian_countries_by_area'
    r = requests.get(url)
except Exception as e:
    print(f"error is {e}        {e.__class__}")
else:
    print('request complete')
soup = BeautifulSoup(r.text, 'html.parser')

My_table = soup.find('table', {'class': 'sortable wikitable sticky-header col2left'} ) 
rows = My_table.findAll('tr')
n = len(rows)
for i in range(1, n-1):
    tag = rows[i].find('a')
    country_names.append(tag.text)

col_text_lists = []
for row in rows:
    l = []
    cols = row.findAll('td')
    for col in cols:
        l.append(col.text)
    col_text_lists.append(l)

for c in col_text_lists:
    if c.__len__() < 1:
        continue
    percent_lists.append(c[2])

# print(country_names,'\n')
# print(percent_lists,'\n')
# print(arae_lists,'\n')

doc = Document()
m, n = len(country_names), 2
table = doc.add_table(rows=m, cols=n)

cell = table.cell(0,0)
p = cell.paragraphs[0]
p.style = 'Heading 1'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Country / dependency')

cell = table.cell(0,1)
p = cell.paragraphs[0]
p.style = 'Heading 1'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('total')

for i in range(1,m):
    cell = table.cell(i, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(country_names[i])

    cell = table.cell(i, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(percent_lists[i].replace('%',''))


table.style = 'Table Grid'
try:
    doc.save('./outs/final5.docx')
except Exception as e:
    print(f"error is {e}        {e.__class__}")
else:
    print('save document complete')
del(doc)