import requests
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document


service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service)

with webdriver.Chrome(service=service) as driver:
    url = 'http://localhost/dashboard/ex03.html'
    driver.get(url)
    h1 = driver.find_element(By.ID,"h01").text
    p1 = driver.find_element(By.ID,"p01").text
    h2 = driver.find_element(By.ID,"h02").text
    p2 = driver.find_element(By.ID,"p02").text

    doc = Document()
    m, n = 2, 2
    table = doc.add_table(rows=m, cols=n)

    cell = table.cell(0,0)
    p = cell.paragraphs[0]
    p.style = 'Heading 1'
    run = p.add_run(h1)

    cell = table.cell(0,1)
    p = cell.paragraphs[0]
    run = p.add_run(p1)

    cell = table.cell(1,0)
    p = cell.paragraphs[0]
    p.style = 'Heading 2'
    run = p.add_run(h2)

    cell = table.cell(1,1)
    p = cell.paragraphs[0]
    run = p.add_run(p2)

    table.style = 'Table Grid'
    try:
        doc.save('./outs/homwork6_1.docx')
    except Exception as e:
        print(f"error is {e}        {e.__class__}")
    else:
        print('save document1 complete')
    del(doc)

with webdriver.Chrome(service=service) as driver:
    url = 'http://localhost/dashboard/ex03.html'  
    driver.get(url)
    link_text = driver.find_element(By.ID, 'h01')
    link_text.click()
    h1 = driver.find_element(By.ID,"h01").text
    p1 = driver.find_element(By.ID,"p01").text
    h2 = driver.find_element(By.ID,"h02").text
    p2 = driver.find_element(By.ID,"p02").text

    doc = Document()
    m, n = 2, 2
    table = doc.add_table(rows=m, cols=n)

    cell = table.cell(0,0)
    p = cell.paragraphs[0]
    p.style = 'Heading 1'
    run = p.add_run(h1)

    cell = table.cell(0,1)
    p = cell.paragraphs[0]
    run = p.add_run(p1)

    cell = table.cell(1,0)
    p = cell.paragraphs[0]
    p.style = 'Heading 2'
    run = p.add_run(h2)

    cell = table.cell(1,1)
    p = cell.paragraphs[0]
    run = p.add_run(p2)

    table.style = 'Table Grid'
    try:
        doc.save('./outs/homwork6_2.docx')
    except Exception as e:
        print(f"error is {e}        {e.__class__}")
    else:
        print('save document2 complete')
    del(doc)
